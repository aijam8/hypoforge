"""Document parsers: PDF, DOCX, TXT, Markdown, LaTeX."""
from __future__ import annotations

import os
import re
from collections import Counter


def parse(path: str, ext: str) -> dict:
    notes: list[str] = []
    if ext == ".pdf":
        text, meta = _read_pdf(path, notes)
    elif ext == ".docx":
        text, meta = _read_docx(path, notes)
    else:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        meta = {}
        notes.append(f"Read plain text ({ext}).")

    text = text.strip()
    words = re.findall(r"[A-Za-z][A-Za-z\-']+", text)
    profile = {
        "char_count": len(text),
        "word_count": len(words),
        **meta,
        "top_terms": _top_terms(words),
        "sections": _guess_sections(text),
    }
    summary = _summarize(text, profile, os.path.basename(path))
    return {
        "kind": "document",
        "filename": os.path.basename(path),
        "summary": summary,
        "profile": profile,
        "dataframe": None,
        "text": text,
        "preview": text[:3000],
        "notes": notes,
    }


def _read_pdf(path: str, notes: list[str]):
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = []
    for pg in reader.pages:
        try:
            pages.append(pg.extract_text() or "")
        except Exception:
            pages.append("")
    notes.append(f"Extracted text from {len(pages)} PDF page(s) via pypdf.")
    meta = {"n_pages": len(pages)}
    info = getattr(reader, "metadata", None)
    if info:
        meta["title"] = str(getattr(info, "title", "") or "")
        meta["author"] = str(getattr(info, "author", "") or "")
    return "\n\n".join(pages), meta


def _read_docx(path: str, notes: list[str]):
    import docx

    d = docx.Document(path)
    text = "\n".join(p.text for p in d.paragraphs)
    notes.append(f"Read DOCX with {len(d.paragraphs)} paragraphs.")
    return text, {"n_paragraphs": len(d.paragraphs)}


_STOP = set("""the a an and or of to in for on with at by from as is are was were be been
being this that these those it its their his her our your my we you they he she i not no
which who whom whose what when where why how can could should would may might will shall
do does did has have had but if then than so such into over under between out up down then
also more most some any all each other into about above below""".split())


def _top_terms(words: list[str], k: int = 25) -> list[dict]:
    counts = Counter(w.lower() for w in words
                     if len(w) > 3 and w.lower() not in _STOP)
    return [{"term": t, "count": c} for t, c in counts.most_common(k)]


def _guess_sections(text: str) -> list[str]:
    heads = []
    for line in text.splitlines():
        s = line.strip()
        if 3 < len(s) < 60 and (s.isupper() or re.match(
                r"^(\d+\.?\s+)?(abstract|introduction|methods?|results?|discussion|"
                r"conclusion|references|background|related work)\b", s, re.I)):
            heads.append(s)
        if len(heads) >= 20:
            break
    return heads


def _summarize(text: str, profile: dict, fname: str) -> str:
    lines = [f"Document '{fname}': {profile.get('word_count', 0)} words"
             + (f", {profile['n_pages']} pages" if profile.get("n_pages") else "") + "."]
    if profile.get("title"):
        lines.append(f"Title metadata: {profile['title']}")
    if profile.get("sections"):
        lines.append("Detected sections: " + ", ".join(profile["sections"][:10]))
    if profile.get("top_terms"):
        lines.append("Most frequent terms: " + ", ".join(
            t["term"] for t in profile["top_terms"][:15]))
    lines.append("Opening excerpt: " + text[:600].replace("\n", " "))
    return "\n".join(lines)
